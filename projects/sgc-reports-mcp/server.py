from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

import markdown as md
import xlsxwriter
from docx import Document
from docxtpl import DocxTemplate
from mcp.server.fastmcp import FastMCP
from pydantic import BaseModel, Field
from weasyprint import HTML

BASE_DIR = Path(__file__).parent
OUT_DIR = BASE_DIR / "out"
OUT_DIR.mkdir(parents=True, exist_ok=True)

mcp = FastMCP("sgc-reports")


class Section(BaseModel):
    heading: str
    body: str


class ReportData(BaseModel):
    title: str
    summary: str
    sections: list[Section] = Field(default_factory=list)
    metrics: dict[str, float | int | str] = Field(default_factory=dict)
    created_at: str = Field(default_factory=lambda: datetime.now().isoformat(timespec="seconds"))


def _safe_name(name: str, ext: str) -> str:
    stem = "".join(c for c in name if c.isalnum() or c in ("-", "_", ".")).strip("._")
    if not stem:
        stem = f"sgc-report-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    if not stem.lower().endswith(ext):
        stem += ext
    return stem


def _report_to_html(report: ReportData) -> str:
    sections_html = "\n".join(
        f"<h2>{s.heading}</h2><div>{md.markdown(s.body)}</div>" for s in report.sections
    )
    metrics_rows = "\n".join(
        f"<tr><td>{k}</td><td>{v}</td></tr>" for k, v in report.metrics.items()
    )
    return f"""
<!doctype html>
<html>
<head>
<meta charset='utf-8'>
<style>
body {{ font-family: Arial, sans-serif; margin: 28px; color: #111; }}
h1 {{ margin-bottom: 0; }}
small {{ color: #666; }}
table {{ border-collapse: collapse; margin-top: 12px; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
tr:nth-child(even) {{ background: #f8f8f8; }}
</style>
</head>
<body>
  <h1>{report.title}</h1>
  <small>Generado: {report.created_at}</small>
  <h2>Resumen</h2>
  <div>{md.markdown(report.summary)}</div>
  <h2>Métricas</h2>
  <table>
    <thead><tr><th>Métrica</th><th>Valor</th></tr></thead>
    <tbody>{metrics_rows}</tbody>
  </table>
  {sections_html}
</body>
</html>
"""


@mcp.tool()
def build_report_data(
    title: str,
    summary: str,
    sections: list[dict[str, str]] | None = None,
    metrics: dict[str, float | int | str] | None = None,
) -> str:
    """Construye el payload estándar del reporte en JSON."""
    payload = ReportData(
        title=title,
        summary=summary,
        sections=[Section(**s) for s in (sections or [])],
        metrics=metrics or {},
    )
    return payload.model_dump_json(indent=2)


@mcp.tool()
def render_docx(data: str, output_name: str, template_path: str | None = None) -> str:
    """Genera .docx desde JSON de ReportData. Usa template docxtpl si se indica."""
    report = ReportData.model_validate_json(data)
    filename = _safe_name(output_name, ".docx")
    out_path = OUT_DIR / filename

    if template_path:
        tpath = Path(template_path).expanduser().resolve()
        if not tpath.exists():
            raise FileNotFoundError(f"Template no encontrado: {tpath}")
        tpl = DocxTemplate(str(tpath))
        context: dict[str, Any] = {
            "title": report.title,
            "summary": report.summary,
            "created_at": report.created_at,
            "metrics": report.metrics,
            "sections": [s.model_dump() for s in report.sections],
        }
        tpl.render(context)
        tpl.save(str(out_path))
    else:
        doc = Document()
        doc.add_heading(report.title, level=1)
        doc.add_paragraph(f"Generado: {report.created_at}")
        doc.add_heading("Resumen", level=2)
        doc.add_paragraph(report.summary)

        if report.metrics:
            doc.add_heading("Métricas", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Métrica"
            table.rows[0].cells[1].text = "Valor"
            for k, v in report.metrics.items():
                row = table.add_row().cells
                row[0].text = str(k)
                row[1].text = str(v)

        for sec in report.sections:
            doc.add_heading(sec.heading, level=2)
            doc.add_paragraph(sec.body)

        doc.save(str(out_path))

    return str(out_path)


@mcp.tool()
def render_xlsx(data: str, output_name: str) -> str:
    """Genera .xlsx desde JSON de ReportData."""
    report = ReportData.model_validate_json(data)
    filename = _safe_name(output_name, ".xlsx")
    out_path = OUT_DIR / filename

    wb = xlsxwriter.Workbook(str(out_path))
    ws_summary = wb.add_worksheet("Resumen")
    ws_sections = wb.add_worksheet("Secciones")

    bold = wb.add_format({"bold": True})

    ws_summary.write("A1", "Título", bold)
    ws_summary.write("B1", report.title)
    ws_summary.write("A2", "Generado", bold)
    ws_summary.write("B2", report.created_at)
    ws_summary.write("A4", "Resumen", bold)
    ws_summary.write("A5", report.summary)

    ws_summary.write("A7", "Métrica", bold)
    ws_summary.write("B7", "Valor", bold)
    r = 7
    for k, v in report.metrics.items():
        ws_summary.write(r, 0, str(k))
        ws_summary.write(r, 1, str(v))
        r += 1

    ws_sections.write("A1", "Heading", bold)
    ws_sections.write("B1", "Body", bold)
    for i, sec in enumerate(report.sections, start=1):
        ws_sections.write(i, 0, sec.heading)
        ws_sections.write(i, 1, sec.body)

    wb.close()
    return str(out_path)


@mcp.tool()
def render_pdf(data: str, output_name: str) -> str:
    """Genera .pdf desde JSON de ReportData usando HTML + WeasyPrint."""
    report = ReportData.model_validate_json(data)
    filename = _safe_name(output_name, ".pdf")
    out_path = OUT_DIR / filename

    html = _report_to_html(report)
    HTML(string=html).write_pdf(str(out_path))
    return str(out_path)


if __name__ == "__main__":
    mcp.run(transport="stdio")
